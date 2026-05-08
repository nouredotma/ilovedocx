from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def apply_edit(file_path: str, instructions: dict) -> bool:
    """Applies structured edit instructions to a .docx file."""
    try:
        doc = Document(file_path)
        action = instructions.get("action")
        params = instructions.get("params", {})

        if action == "replace_text":
            find_text = params.get("find")
            replace_with = params.get("replace")
            for p in doc.paragraphs:
                if find_text in p.text:
                    for run in p.runs:
                        if find_text in run.text:
                            run.text = run.text.replace(find_text, replace_with)

        elif action == "change_font":
            target = params.get("target")
            for p in doc.paragraphs:
                if target in p.text:
                    for run in p.runs:
                        if params.get("font_name"):
                            run.font.name = params["font_name"]
                        if params.get("font_size"):
                            run.font.size = Pt(params["font_size"])
                        if "bold" in params:
                            run.font.bold = params["bold"]
                        if "italic" in params:
                            run.font.italic = params["italic"]
                    break

        elif action == "edit_header":
            section = doc.sections[0]
            header = section.header
            # Clear existing content
            for p in header.paragraphs:
                for run in p.runs:
                    run.text = ""
            # Add new text to the first paragraph
            if not header.paragraphs:
                header.add_paragraph(params.get("new_text", ""))
            else:
                header.paragraphs[0].text = params.get("new_text", "")

        elif action == "edit_footer":
            section = doc.sections[0]
            footer = section.footer
            for p in footer.paragraphs:
                for run in p.runs:
                    run.text = ""
            if not footer.paragraphs:
                footer.add_paragraph(params.get("new_text", ""))
            else:
                footer.paragraphs[0].text = params.get("new_text", "")

        elif action == "remove_image":
            index = params.get("image_index", 0)
            inline_shapes = doc.inline_shapes
            if 0 <= index < len(inline_shapes):
                shape = inline_shapes[index]
                element = shape._inline
                element.getparent().remove(element)

        elif action == "insert_paragraph":
            after_text = params.get("after_text")
            new_text = params.get("new_text")
            style = params.get("style", "Normal")
            for p in doc.paragraphs:
                if after_text in p.text:
                    new_p = doc.add_paragraph(new_text, style=style)
                    # Move new_p after p
                    p._element.addnext(new_p._element)
                    break

        elif action == "delete_paragraph":
            target = params.get("target")
            for p in doc.paragraphs:
                if target in p.text:
                    p._element.getparent().remove(p._element)
                    break

        elif action == "change_alignment":
            target = params.get("target")
            align_str = params.get("alignment", "LEFT").upper()
            align_map = {
                "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
                "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
                "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
                "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY
            }
            target_align = align_map.get(align_str, WD_ALIGN_PARAGRAPH.LEFT)
            for p in doc.paragraphs:
                if target in p.text:
                    p.alignment = target_align
                    break

        elif action == "answer_only":
            return True

        # Save changes for all edit actions
        if action != "answer_only":
            doc.save(file_path)
        
        return True
    except Exception as e:
        print(f"Editor error: {e}")
        return False
