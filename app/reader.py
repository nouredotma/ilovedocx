from docx import Document

def extract_content(file_path: str) -> str:
    """Extracts text content from header, body, footer, and tables of a .docx file."""
    try:
        doc = Document(file_path)
        content = []

        # Header
        header_text = []
        for section in doc.sections:
            for p in section.header.paragraphs:
                if p.text.strip():
                    header_text.append(p.text)
        if header_text:
            content.append("[HEADER]\n" + "\n".join(header_text))

        # Body
        body_text = [p.text for p in doc.paragraphs if p.text.strip()]
        if body_text:
            content.append("[BODY]\n" + "\n".join(body_text))

        # Footer
        footer_text = []
        for section in doc.sections:
            for p in section.footer.paragraphs:
                if p.text.strip():
                    footer_text.append(p.text)
        if footer_text:
            content.append("[FOOTER]\n" + "\n".join(footer_text))

        # Tables
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_content = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_content:
                    table_text.append(" | ".join(row_content))
        if table_text:
            content.append("[TABLES]\n" + "\n".join(table_text))

        return "\n\n".join(content) if content else "Document is empty."
    except Exception as e:
        print(f"Error reading docx: {e}")
        return "Could not read document content"
